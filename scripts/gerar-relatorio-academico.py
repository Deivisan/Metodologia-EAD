#!/usr/bin/env python3
"""Gera relatorio academico MinhaFila .docx com formatacao ABNT."""
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "/home/deivi/Projetos/Metodologia-EAD/disciplinas/gcetens843-projeto-algoritmo-i/atividades/trabalhos/projeto-algoritmos/relatorio-academico-minhafila.docx"
LOGO = "/home/deivi/Projetos/Metodologia-EAD/templates/logo-ufrb-20-anos.png"

doc = Document()

# ─── Configurar margens ABNT ───
for section in doc.sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

# ─── Estilos base ───
style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(12)
style_normal.paragraph_format.line_spacing = 1.5
style_normal.paragraph_format.space_after = Pt(0)
style_normal.paragraph_format.space_before = Pt(0)

# Heading styles
for level in [1, 2, 3]:
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Times New Roman'
    h.font.color.rgb = RGBColor(0, 0, 0)
    h.font.bold = True
    h.paragraph_format.line_spacing = 1.5
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    if level == 1:
        h.font.size = Pt(14)
    elif level == 2:
        h.font.size = Pt(13)
    else:
        h.font.size = Pt(12)

def add_para(text, bold=False, italic=False, size=12, align=None, space_after=0, space_before=0, style=None):
    """Adiciona paragrafo formatado."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p

def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

# ════════════════════════════════════════
# CAPA
# ════════════════════════════════════════

# Logo
if os.path.exists(LOGO):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(LOGO, width=Cm(12))
    p.paragraph_format.space_after = Pt(24)

add_para('UNIVERSIDADE FEDERAL DO RECÔNCAVO DA BAHIA', bold=True, size=13, align='center', space_after=2)
add_para('CENTRO DE CIÊNCIA E TECNOLOGIA — CETENS', bold=True, size=12, align='center', space_after=2)
add_para('BACHARELADO EM SISTEMAS DE INFORMAÇÃO (EAD)', bold=True, size=12, align='center', space_after=24)

# Linha horizontal
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('_' * 60)
run.font.name = 'Times New Roman'
run.font.size = Pt(10)
p.paragraph_format.space_after = Pt(36)

# Titulo do trabalho
add_para('MinhaFila:', bold=False, size=18, align='center', space_after=6)
add_para('Sistema de Gerenciamento Inteligente de Filas Presenciais', bold=False, size=18, align='center', space_after=4)
add_para('em Serviços de Saúde com Foco em Acessibilidade', bold=False, size=18, align='center', space_after=36)

add_para('GCETENS843 — Projeto Algoritmo I', bold=True, size=13, align='center', space_after=6)
add_para('Prof. Luis Paulo Morais Conceição', size=12, align='center', space_after=6)
add_para('Semestre 2026.1', size=12, align='center', space_after=6)
add_para('29 de junho de 2026', size=12, align='center', space_after=36)

# Equipe
add_para('Equipe:', bold=True, size=12, align='center', space_after=4)
membros = [
    'Deivison de Lima Santana', 'Ausiane', 'Nubia de Asiká', 'Rios',
    'Artur Campos', 'Marcos Vinícius', 'Uélinton Cardoso Santana', 'Wallace'
]
for nome in membros:
    add_para(f'• {nome}', size=11, align='center', space_after=1)

# Pular para segunda pagina
doc.add_page_break()

# ════════════════════════════════════════
# RESUMO
# ════════════════════════════════════════

add_heading('Resumo', level=1)

add_para(
    'O presente trabalho apresenta o MinhaFila, um sistema de gerenciamento inteligente de filas '
    'presenciais para serviços de saúde, desenvolvido com foco em transparência, acessibilidade e '
    'inclusão. A solução propõe o acompanhamento de filas em tempo real por meio de dispositivo '
    'móvel, com geração de senha virtual, notificação antecipada de atendimento e recursos de '
    'acessibilidade para pessoas com deficiência auditiva, visual, Transtorno do Espectro Autista '
    '(TEA), idosos e gestantes. O sistema visa reduzir o tempo de permanência física dos usuários '
    'nas unidades de saúde, oferecer previsibilidade sobre o atendimento e garantir que pessoas '
    'com necessidades especiais não sejam excluídas do processo. O trabalho inclui a fundamentação '
    'do problema real, análise de soluções existentes, fluxograma do sistema, protótipo de telas '
    'e considerações sobre a viabilidade técnica da proposta.',
    size=12, align='justify', space_after=12
)

add_para(
    'Palavras-chave: filas presenciais; saúde; acessibilidade; inclusão; sistema de informação.',
    italic=True, size=11, align='justify', space_after=18
)

# ════════════════════════════════════════
# 1. INTRODUCAO
# ════════════════════════════════════════

add_heading('1. Introdução', level=1)

add_para(
    'A transformação digital tem promovido mudanças significativas em diversos setores da sociedade, '
    'incluindo o setor da saúde. O avanço das tecnologias digitais, dos sistemas de informação em '
    'saúde e das soluções móveis tem contribuído para a melhoria da gestão dos serviços, da '
    'comunicação entre profissionais e usuários e da qualidade da assistência prestada '
    '(Aggarwal et al., 2022; Kessel et al., 2022).',
    size=12, align='justify', space_after=6
)

add_para(
    'No entanto, um problema persiste em hospitais, clínicas, unidades básicas de saúde (UBSs) e '
    'Unidades de Pronto Atendimento (UPAs): a demora nas filas presenciais de atendimento e a falta '
    'de informações claras sobre a posição do usuário na fila ou sobre o tempo estimado para o '
    'atendimento.',
    size=12, align='justify', space_after=6
)

add_para(
    'Este problema é agravado para pessoas com deficiência. Indivíduos com deficiência auditiva podem '
    'não ouvir chamadas por alto-falante; pessoas com baixa visão podem não ler painéis com códigos '
    'opacos; e pessoas com Transtorno do Espectro Autista (TEA) podem sofrer com o excesso de '
    'estímulos sensoriais e a imprevisibilidade do ambiente de espera '
    '(Clemente et al., 2022; Farias et al., 2023; Ha et al., 2023).',
    size=12, align='justify', space_after=6
)

add_para(
    'Diante desse cenário, o presente projeto propõe o MinhaFila, um sistema de gerenciamento '
    'inteligente de filas presenciais para serviços de saúde, desenvolvido com transparência radical '
    'e acessibilidade como pilares centrais. A pergunta que orienta esta pesquisa é: como reduzir o '
    'tempo de permanência física dos usuários nas unidades de saúde e aumentar a previsibilidade do '
    'atendimento por meio de uma solução digital acessível e inclusiva?',
    size=12, align='justify', space_after=12
)

# ════════════════════════════════════════
# 2. PROBLEMA REAL
# ════════════════════════════════════════

add_heading('2. Problema Real', level=1)

add_heading('2.1 Contextualização', level=2)

add_para(
    'O fluxo atual em serviços de saúde de porta aberta segue, de maneira geral, as seguintes '
    'etapas: o usuário chega à unidade, dirige-se ao guichê de recepção, recebe uma senha física '
    '(baseada na ordem de chegada) e aguarda na sala de espera até ser chamado por alto-falante '
    'ou painel eletrônico. A chamada é geralmente feita por meio de códigos opacos como "AB-22" '
    'ou "H-27", sem qualquer informação sobre a posição do usuário na fila ou o tempo estimado '
    'de espera.',
    size=12, align='justify', space_after=6
)

add_para(
    'Estudos demonstram que tempos prolongados de espera estão associados à redução da satisfação '
    'dos usuários e à percepção negativa da qualidade dos serviços de saúde (Aggarwal et al., 2022). '
    'A espera prolongada também está associada à aglomeração de pacientes, falhas na comunicação, '
    'aumento do estresse, absenteísmo em consultas e redução da eficiência operacional dos serviços '
    '(Thompson et al., 2023).',
    size=12, align='justify', space_after=6
)

add_heading('2.2 O Problema da Acessibilidade', level=2)

add_para(
    'Pessoas com deficiência enfrentam barreiras adicionais significativas. Indivíduos com deficiência '
    'auditiva podem não perceber chamadas realizadas exclusivamente por áudio. Pessoas com baixa '
    'visão ou deficiência visual encontram dificuldades para ler painéis com caracteres pequenos '
    'ou em cores de baixo contraste. Pessoas com TEA podem ser sobrecarregadas sensorialmente pelo '
    'ambiente de sala de espera, com ruídos, luzes e aglomeração. Idosos e gestantes, por sua vez, '
    'têm dificuldade em permanecer em pé por longos períodos.',
    size=12, align='justify', space_after=6
)

add_para(
    'Embora existam dispositivos legais que assegurem atendimento prioritário a esses grupos, a '
    'acessibilidade física e digital dos serviços de saúde ainda é insuficiente '
    '(Farias et al., 2023; Jonsson et al., 2023).',
    size=12, align='justify', space_after=6
)

add_heading('2.3 Pergunta Central', level=2)

add_para(
    'A questão que orienta este projeto é:',
    size=12, align='justify', space_after=6
)

# Citação em bloco (recuada)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.left_indent = Cm(4)
p.paragraph_format.right_indent = Cm(2)
p.paragraph_format.line_spacing = 1.0
p.paragraph_format.space_after = Pt(12)
run = p.add_run(
    'Como reduzir o tempo de permanência física dos usuários nas unidades de saúde e aumentar '
    'a previsibilidade do atendimento por meio de uma solução digital acessível e inclusiva?'
)
run.font.name = 'Times New Roman'
run.font.size = Pt(10)
run.italic = True

# ════════════════════════════════════════
# 3. SOLUÇÕES EXISTENTES
# ════════════════════════════════════════

add_heading('3. Soluções Existentes', level=1)

add_para(
    'A pesquisa por soluções existentes no mercado foi realizada pela equipe com o objetivo de '
    'identificar sistemas concorrentes ou complementares que já atuam no segmento de gerenciamento '
    'de filas em serviços de saúde. Os resultados desta pesquisa estão documentados em arquivo '
    'complementar e incluem aplicativos e plataformas que oferecem funcionalidades como senha '
    'digital, acompanhamento remoto e painéis de chamada.',
    size=12, align='justify', space_after=6
)

add_para(
    'A análise dessas soluções existentes indicou que, embora haja sistemas de gerenciamento de '
    'filas disponíveis, a maioria não incorpora a acessibilidade como um pilar central do design, '
    'limitando-se a oferecer recursos básicos que não atendem adequadamente pessoas com deficiência '
    'auditiva, visual ou TEA.',
    size=12, align='justify', space_after=12
)

# ════════════════════════════════════════
# 4. PROPOSTA DE SOLUCAO
# ════════════════════════════════════════

add_heading('4. Proposta de Solução: MinhaFila', level=1)

add_heading('4.1 Visão Geral', level=2)

add_para(
    'O MinhaFila é um sistema de acompanhamento de filas em tempo real para serviços de saúde, '
    'com foco em transparência e acessibilidade. A solução permite que o paciente: (a) realize um '
    'cadastro rápido ao chegar à unidade de saúde; (b) receba uma senha virtual diretamente em '
    'seu dispositivo móvel; (c) acompanhe em tempo real sua posição na fila e o tempo estimado '
    'de espera; (d) receba notificações antecipadas quando o atendimento estiver próximo; '
    '(e) visualize informações claras sobre as pessoas à sua frente (quantidade e classificação '
    'de risco, sem identificação nominal); e (f) utilize recursos de acessibilidade como alto '
    'contraste, texto ampliado, modo silencioso para surdos e modo baixo estímulo para TEA.',
    size=12, align='justify', space_after=6
)

add_heading('4.2 Funcionalidades Principais', level=2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_after = Pt(6)
run = p.add_run('Para o paciente:')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.bold = True

for item in [
    'Cadastro rápido com identificação de necessidades especiais',
    'Senha virtual e posição na fila em tempo real',
    'Tempo estimado de atendimento',
    'Notificação push quando faltarem poucas pessoas',
    'Modos de acessibilidade adaptáveis (alto contraste, texto ampliado, modo silencioso, modo TEA)',
]:
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(item)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5

add_para('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_after = Pt(6)
run = p.add_run('Para a equipe de saúde:')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.bold = True

for item in [
    'Painel interno com lista ordenada por ordem de chegada e classificação de risco',
    'Botão para chamar o próximo paciente',
    'Identificação de pacientes com necessidades especiais',
    'Histórico de atendimentos realizados',
]:
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(item)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5

add_heading('4.3 Diferenciais', level=2)

add_para(
    'O MinhaFila se diferencia das soluções existentes por quatro aspectos principais. '
    'Primeiro, a acessibilidade como pilar central: modos específicos para pessoas com deficiência '
    'auditiva, visual, TEA, idosos e gestantes. Segundo, a transparência radical: exibição da '
    'posição real e cores de risco das pessoas à frente, sem identificação nominal. Terceiro, '
    'a mobilidade: possibilidade de sair do local e ser notificado quando o atendimento estiver '
    'próximo. Quarto, a redução do tempo de permanência física: a previsibilidade permite que o '
    'usuário planeje melhor sua espera.',
    size=12, align='justify', space_after=12
)

# ════════════════════════════════════════
# 5. FLUXOGRAMA
# ════════════════════════════════════════

add_heading('5. Fluxograma do Sistema', level=1)

add_para(
    'O fluxograma do sistema descreve o percurso completo do usuário no MinhaFila, desde a '
    'entrada até o atendimento, incluindo a lógica de prioridade e notificação. O fluxo é '
    'composto pelas seguintes etapas:',
    size=12, align='justify', space_after=6
)

etapas = [
    ('1. Entrada no Sistema', 'O usuário acessa o aplicativo (web ou mobile). Se já possui cadastro, realiza login; caso contrário, faz um cadastro rápido informando dados pessoais e necessidades especiais.'),
    ('2. Solicitação de Atendimento', 'O usuário seleciona a clínica ou unidade de saúde e o tipo de atendimento desejado (consulta comum, exame, atendimento prioritário, retorno ou triagem).'),
    ('3. Decisão de Prioridade', 'O sistema verifica se o usuário se enquadra em alguma categoria de prioridade legal (idoso, gestante, pessoa com deficiência, TEA). Em caso positivo, é direcionado para a fila prioritária; caso contrário, para a fila convencional. Ambas as filas seguem o mesmo fluxo de atendimento.'),
    ('4. Geração de Senha', 'O sistema gera uma senha única, registra o horário de entrada, posiciona o usuário no final da fila correspondente e calcula a previsão inicial de atendimento com base no tempo médio histórico.'),
    ('5. Chamada e Notificação', 'Quando o usuário atinge o topo da fila, é chamado para atendimento. O segundo usuário da fila recebe uma notificação antecipada para se dirigir ao local, reduzindo o tempo de espera entre chamadas.'),
    ('6. Atendimento e Finalização', 'Após o atendimento, o sistema exclui o usuário da fila, que é atualizada automaticamente. O próximo usuário no topo é então chamado.'),
]

for titulo, descricao in etapas:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(titulo)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True

    add_para(descricao, size=12, align='justify', space_after=6)

add_para(
    'O fluxograma completo pode ser visualizado no arquivo complementar fluxograma.pdf '
    'anexo a este relatório.',
    size=12, align='justify', space_after=12
)

# ════════════════════════════════════════
# 6. PROTOTIPO
# ════════════════════════════════════════

add_heading('6. Protótipo de Telas', level=1)

add_para(
    'O protótipo do MinhaFila foi desenvolvido com foco na experiência do usuário e nos princípios '
    'de acessibilidade. As telas principais do fluxo do paciente incluem: tela de splash com a '
    'identidade visual do sistema; onboarding com apresentação das funcionalidades principais; '
    'home com opções de clínicas e atendimentos disponíveis; cadastro com formulário para dados '
    'pessoais e necessidades especiais; senha virtual com exibição da senha gerada e posição na '
    'fila; acompanhamento com posição, tempo estimado e notificações; e notificação de chamada '
    'com alerta visual e sonoro quando o atendimento estiver próximo.',
    size=12, align='justify', space_after=6
)

add_para(
    'O painel da clínica, por sua vez, conta com lista de pacientes ordenada por ordem de chegada '
    'com indicação de prioridade, botão para chamar o próximo paciente e histórico de atendimentos '
    'realizados.',
    size=12, align='justify', space_after=6
)

add_para(
    'O protótipo interativo pode ser acessado em formato HTML, anexo a este relatório.',
    size=12, align='justify', space_after=12
)

# ════════════════════════════════════════
# 7. CONSIDERACOES FINAIS
# ════════════════════════════════════════

add_heading('7. Considerações Finais', level=1)

add_para(
    'O MinhaFila apresenta-se como uma solução viável e necessária para o problema das filas '
    'presenciais em serviços de saúde, combinando tecnologia móvel, transparência e acessibilidade '
    'em um único sistema. A abordagem proposta responde diretamente à pergunta central do projeto, '
    'oferecendo meios para reduzir o tempo de permanência física dos usuários nas unidades de '
    'saúde e aumentar a previsibilidade do atendimento.',
    size=12, align='justify', space_after=6
)

add_para(
    'A pesquisa realizada pela equipe evidenciou que as soluções existentes no mercado não '
    'incorporam a acessibilidade como pilar central do design, deixando desassistidas pessoas com '
    'deficiência auditiva, visual e TEA. O MinhaFila preenche essa lacuna ao oferecer recursos '
    'específicos para cada tipo de necessidade.',
    size=12, align='justify', space_after=6
)

add_para(
    'Como trabalhos futuros, sugere-se o desenvolvimento de integração com sistemas oficiais de '
    'saúde, como o Meu SUS Digital e a Rede Nacional de Dados em Saúde (RNDS), bem como a '
    'implementação do código computacional em linguagem C para validação dos algoritmos de '
    'gerenciamento de filas.',
    size=12, align='justify', space_after=12
)

# ════════════════════════════════════════
# REFERENCIAS
# ════════════════════════════════════════

add_heading('Referências', level=1)

referencias = [
    'AGGARWAL, Ravi et al. Defining the enablers and barriers to the implementation of large-scale, health care-related mobile technology: qualitative case study in a tertiary hospital setting. JMIR mHealth and uHealth, Toronto, v. 10, n. 2, 2022.',
    'CHIANCA, Tânia Couto Machado et al. Tempos de espera para atendimento usando sistema de triagem de Manchester em um hospital de urgência. Reme: Revista Mineira de Enfermagem, v. 20, 2016.',
    'CLEMENTE, Karina Aparecida Padilha et al. Barreiras ao acesso das pessoas com deficiência aos serviços de saúde: uma revisão de escopo. Revista de Saúde Pública, v. 56, n. 64, 2022.',
    'FARIAS, Tássia Mayra Oliveira et al. O estreito acesso das Pessoas com Deficiência aos serviços de saúde em uma capital nordestina. Ciência & Saúde Coletiva, v. 28, n. 5, 2023.',
    'GIANNOTTI, Elaine Maria; LOUVISON, Marília; CHIORO, Arthur. Listas de espera na atenção ambulatorial especializada: reflexões sobre um conceito crítico para o Sistema Único de Saúde. Cadernos de Saúde Pública, v. 41, n. 6, 2025.',
    'HA, Sandeul et al. Digital health equity and tailored health care service for people with disability: user-centered design and usability study. Journal of Medical Internet Research, v. 25, 2023.',
    'JONSSON, Marika et al. Development and evaluation of eHealth services regarding accessibility: scoping literature review. Journal of Medical Internet Research, v. 25, 2023.',
    'KESSEL, Robin Van et al. Digital health paradox: international policy perspectives to address increased health inequalities for people living with disabilities. Journal of Medical Internet Research, Toronto, v. 24, n. 2, 2022.',
    'MELO, Daiane Celestino et al. Acessibilidade aos serviços de saúde e posição dos usuários no espaço social em Salvador, Bahia, 2006: um estudo transversal. Epidemiologia e Serviços de Saúde, v. 30, n. 2, 2021.',
    'THOMPSON, Yee Lam Elim et al. Evaluation of wait time saving effectiveness of triage algorithms. arXiv:2303.07050, v. 1, 2023.',
]

for ref in referencias:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)

# ─── Salvar ───
doc.save(OUTPUT)
print(f'✅ Documento salvo em: {OUTPUT}')
