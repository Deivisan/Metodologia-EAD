#!/usr/bin/env python3
"""Aplica molde institucional UFRB a arquivos .docx:
   - Logo UFRB 20 anos no cabeçalho
   - Cores institucionais (azul UFRB, dourado)
   - Times New Roman, espaçamento 1.5
   - Cabeçalho texto padrão
"""

import sys
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Cores institucionais UFRB
AZUL_UFRB = RGBColor(0x00, 0x3D, 0x6B)     # azul escuro
AZUL_CLARO = RGBColor(0x00, 0x6D, 0xAA)     # azul medio
DOURADO = RGBColor(0xC8, 0x9B, 0x3C)        # dourado
CINZA = RGBColor(0x50, 0x50, 0x50)          # texto secundario
PRETO = RGBColor(0x00, 0x00, 0x00)

LOGO_PATH = os.path.join(os.path.dirname(__file__), '..', 'templates', 'logo-ufrb-20-anos.png')


def adicionar_cabecalho_ufrb(doc, title="Documento"):
    """Adiciona logo e texto institucional no cabecalho + primeira pagina."""

    # === CABECALHO (todas as paginas) ===
    section = doc.sections[0]
    section.top_margin = Inches(1.2)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    header = section.header
    header.is_linked_to_previous = False

    # Tabela no cabecalho: logo | texto
    tbl = header.add_table(rows=1, cols=2, width=Inches(6.5))

    # Celula 1: Logo
    cell_logo = tbl.cell(0, 0)
    cell_logo.width = Inches(1.4)
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if os.path.exists(LOGO_PATH):
        run_logo = p_logo.add_run()
        run_logo.add_picture(LOGO_PATH, width=Inches(1.3))

    # Celula 2: Texto institucional
    cell_texto = tbl.cell(0, 1)
    cell_texto.width = Inches(5.0)
    p_inst = cell_texto.paragraphs[0]
    p_inst.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run_ufrb = p_inst.add_run("UNIVERSIDADE FEDERAL DO RECÔNCAVO DA BAHIA — UFRB")
    run_ufrb.font.size = Pt(8)
    run_ufrb.font.name = "Times New Roman"
    run_ufrb.bold = True
    run_ufrb.font.color.rgb = AZUL_UFRB

    p_inst2 = cell_texto.add_paragraph()
    p_inst2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r2 = p_inst2.add_run("CENTRO DE CIÊNCIA E TECNOLOGIA — CETENS")
    r2.font.size = Pt(7.5)
    r2.font.name = "Times New Roman"
    r2.bold = True
    r2.font.color.rgb = AZUL_CLARO

    p_inst3 = cell_texto.add_paragraph()
    p_inst3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r3 = p_inst3.add_run("BACHARELADO EM SISTEMAS DE INFORMAÇÃO (EAD)")
    r3.font.size = Pt(7)
    r3.font.name = "Times New Roman"
    r3.font.color.rgb = CINZA

    # Linha dourada abaixo
    p_line = cell_texto.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_line = p_line.add_run("─" * 80)
    r_line.font.size = Pt(4)
    r_line.font.color.rgb = DOURADO

    # Remover bordas da tabela do cabecalho
    for cell in tbl.row_cells(0):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            '<w:tcBorders %s>'
            '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '</w:tcBorders>' % nsdecls('w')
        )
        tcPr.append(tcBorders)


def aplicar_molde(docx_path):
    """Abre .docx, aplica molde UFRB e salva."""
    print(f"  → {os.path.basename(docx_path)}...", end=" ")

    doc = Document(docx_path)

    # Adicionar cabecalho
    adicionar_cabecalho_ufrb(doc, title=os.path.basename(docx_path))

    # Ajustar estilos padrao
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.font.color.rgb = PRETO
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)

    # Ajustar headings (se existirem no documento)
    for hname, size, color in [
        ('Heading 1', Pt(14), AZUL_UFRB),
        ('Heading 2', Pt(13), AZUL_CLARO),
        ('Heading 3', Pt(12), AZUL_UFRB),
    ]:
        try:
            h = doc.styles[hname]
            h.font.name = 'Times New Roman'
            h.font.size = size
            h.font.bold = True
            h.font.color.rgb = color
            h.paragraph_format.line_spacing = 1.5
        except KeyError:
            pass

    # Rodape
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p_footer = footer.paragraphs[0]
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_footer = p_footer.add_run(f"MinhaFila — GCETENS843 — 2026.1")
    r_footer.font.size = Pt(7)
    r_footer.font.name = 'Times New Roman'
    r_footer.font.color.rgb = CINZA
    r_footer.italic = True

    doc.save(docx_path)
    print("✅ OK")


def main():
    base = os.path.join(
        os.path.dirname(__file__),
        '..',
        'disciplinas',
        'gcetens843-projeto-algoritmo-i',
        'atividades',
        'trabalhos',
        'projeto-algoritmos'
    )

    arquivos = [
        'problema-real.docx',
        'proposta.docx',
        'fluxograma.docx',
        'Index-Geral.docx',
        'relatorio-academico-minhafila.docx',
    ]

    if not os.path.exists(LOGO_PATH):
        print(f"❌ Logo nao encontrado em: {LOGO_PATH}")
        sys.exit(1)

    print(f"🎨 Aplicando molde UFRB ({len(arquivos)} arquivos)...\n")

    for arq in arquivos:
        caminho = os.path.join(base, arq)
        if not os.path.exists(caminho):
            print(f"  ⚠ {arq} nao encontrado, pulando")
            continue
        aplicar_molde(caminho)

    print(f"\n✅ Todos os arquivos atualizados com o molde UFRB!")


if __name__ == '__main__':
    main()
