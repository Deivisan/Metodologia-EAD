#!/usr/bin/env python3
"""Gera Index-Geral.docx com molde UFRB a partir do Index-Geral.md."""

import os
import subprocess
import sys

LOGO_PATH = os.path.abspath(os.path.join(
    os.path.dirname(__file__), '..', 'templates', 'logo-ufrb-20-anos.png'
))

BASE = os.path.abspath(os.path.join(
    os.path.dirname(__file__), '..',
    'disciplinas', 'gcetens843-projeto-algoritmo-i',
    'atividades', 'trabalhos', 'projeto-algoritmos'
))

AZUL_UFRB_HEX = '003D6B'
AZUL_CLARO_HEX = '006DAA'
DOURADO_HEX = 'C89B3C'
CINZA_HEX = '505050'


def gerar_docx_via_pandoc():
    """Gera .docx a partir do .md usando pandoc."""
    md_path = os.path.join(BASE, 'Index-Geral.md')
    docx_path = os.path.join(BASE, 'Index-Geral.docx')

    if not os.path.exists(md_path):
        print(f"❌ Index-Geral.md não encontrado em {md_path}")
        return None

    # Gera .docx com pandoc (da raiz do repo para resolver caminhos relativos)
    repo_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
    cmd = [
        'pandoc', md_path,
        '-o', docx_path,
        '--from', 'markdown',
        '--to', 'docx',
        '--resource-path', os.path.dirname(md_path),
    ]
    result = subprocess.run(cmd, capture_output=True, text=True, cwd=repo_root)
    if result.returncode != 0:
        print(f"  ❌ Erro pandoc: {result.stderr}")
        return None

    size = os.path.getsize(docx_path)
    print(f"  ✓ Index-Geral.docx gerado via pandoc ({size/1024:.0f} KB)")
    return docx_path


def aplicar_molde_ufrb(docx_path):
    """Abre .docx, insere capa com logo UFRB, aplica estilos e rodapé."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    doc = Document(docx_path)

    # ========== CAPA ==========
    capa_tbl = doc.add_table(rows=1, cols=2)
    capa_tbl.autofit = False
    capa_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Remover bordas da tabela
    tbl_pr = capa_tbl._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        capa_tbl._tbl.insert(0, tbl_pr)

    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tbl_pr.append(borders)

    # Célula esquerda: logo
    cell_logo = capa_tbl.cell(0, 0)
    cell_logo.width = Inches(1.4)
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists(LOGO_PATH):
        run_logo = p_logo.add_run()
        run_logo.add_picture(LOGO_PATH, width=Inches(1.25))

    # Célula direita: texto institucional
    cell_texto = capa_tbl.cell(0, 1)
    cell_texto.width = Inches(5.0)

    def add_texto(parag, texto, size, bold, color_hex):
        run = parag.add_run(texto)
        run.font.size = Pt(size)
        run.font.name = 'Times New Roman'
        run.bold = bold
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16),
            int(color_hex[2:4], 16),
            int(color_hex[4:6], 16)
        )
        return run

    p1 = cell_texto.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_texto(p1, "UNIVERSIDADE FEDERAL DO RECÔNCAVO DA BAHIA — UFRB", 8, True, AZUL_UFRB_HEX)

    p2 = cell_texto.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_texto(p2, "CENTRO DE CIÊNCIA E TECNOLOGIA — CETENS", 7.5, True, AZUL_CLARO_HEX)

    p3 = cell_texto.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_texto(p3, "BACHARELADO EM SISTEMAS DE INFORMAÇÃO (EAD)", 7, False, CINZA_HEX)

    p4 = cell_texto.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_texto(p4, "─" * 80, 4, False, DOURADO_HEX)

    # Mover a tabela para o início do documento
    tbl_element = capa_tbl._tbl
    first_para = doc.paragraphs[0]._element
    first_para.addprevious(tbl_element)

    # ========== RODAPÉ ==========
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()
    for tbl in footer.tables:
        tbl._element.getparent().remove(tbl._element)

    p_footer = footer.paragraphs[0]
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_footer = p_footer.add_run("MinhaFila — GCETENS843 — 2026.1")
    r_footer.font.size = Pt(7)
    r_footer.font.name = 'Times New Roman'
    r_footer.font.color.rgb = RGBColor(
        int(CINZA_HEX[0:2], 16), int(CINZA_HEX[2:4], 16), int(CINZA_HEX[4:6], 16)
    )
    r_footer.italic = True

    # ========== ESTILOS DO DOCUMENTO ==========
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    for hname, size, chex in [
        ('Heading 1', 14, AZUL_UFRB_HEX),
        ('Heading 2', 13, AZUL_CLARO_HEX),
        ('Heading 3', 12, AZUL_UFRB_HEX),
    ]:
        try:
            h = doc.styles[hname]
            h.font.name = 'Times New Roman'
            h.font.size = Pt(size)
            h.font.bold = True
            h.font.color.rgb = RGBColor(
                int(chex[0:2], 16), int(chex[2:4], 16), int(chex[4:6], 16)
            )
            h.paragraph_format.line_spacing = 1.5
        except KeyError:
            pass

    # ========== TÍTULO CENTRAL ==========
    # Insere título "Index Geral — MinhaFila" centralizado após a capa
    # O primeiro heading 1 já existe, vamos deixá-lo como está

    doc.save(docx_path)
    print(f"  ✓ molde UFRB aplicado em Index-Geral.docx")


def main():
    print("📄 Regenerando Index-Geral.docx com molde UFRB...\n")

    if not os.path.exists(LOGO_PATH):
        print(f"❌ Logo não encontrado: {LOGO_PATH}")
        sys.exit(1)

    docx = gerar_docx_via_pandoc()
    if docx:
        aplicar_molde_ufrb(docx)
        size = os.path.getsize(docx)
        print(f"\n✅ Index-Geral.docx gerado com sucesso ({size/1024:.0f} KB)")
    else:
        print("\n❌ Falha ao gerar Index-Geral.docx")
        sys.exit(1)


if __name__ == '__main__':
    main()
