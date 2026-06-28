#!/usr/bin/env python3
"""Aplica molde UFRB a arquivos .docx do zero:
   1. Regenera .docx via pandoc a partir do .md (limpo)
   2. Aplica header com logo UFRB (substituindo header existente)
"""

import sys
import os
import subprocess

LOGO_PATH = os.path.abspath(os.path.join(
    os.path.dirname(__file__), '..', 'templates', 'logo-ufrb-20-anos.png'
))

BASE = os.path.abspath(os.path.join(
    os.path.dirname(__file__), '..',
    'disciplinas', 'gcetens843-projeto-algoritmo-i',
    'atividades', 'trabalhos', 'projeto-algoritmos'
))

ARQUIVOS = [
    'problema-real',
    'proposta',
    'fluxograma',
]

AZUL_UFRB_HEX = '003D6B'
AZUL_CLARO_HEX = '006DAA'
DOURADO_HEX = 'C89B3C'
CINZA_HEX = '505050'


def gerar_docx_via_pandoc(nome):
    """Gera .docx limpo a partir do .md usando pandoc."""
    md_path = os.path.join(BASE, f'{nome}.md')
    docx_path = os.path.join(BASE, f'{nome}.docx')

    if not os.path.exists(md_path):
        print(f"  ⚠ {nome}.md não encontrado, pulando")
        return None

    # Gera .docx com pandoc
    cmd = [
        'pandoc', md_path,
        '-o', docx_path,
        '--from', 'markdown',
        '--to', 'docx',
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"  ❌ Erro pandoc {nome}: {result.stderr}")
        return None

    print(f"  ✓ {nome}.docx gerado via pandoc")
    return docx_path


def aplicar_header_ufrb(docx_path):
    """Abre .docx e substitui/adiciona header com logo UFRB + texto."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    import copy

    doc = Document(docx_path)

    # ========== HEADER ==========
    section = doc.sections[0]

    # Remove header existente (se houver) — apaga conteudo do paragrafo padrao
    header = section.header
    header.is_linked_to_previous = False

    # Limpa TODOS os elementos do header (remove paragrafos antigos)
    for p in header.paragraphs:
        p.clear()

    # Remove tabelas antigas do header
    for tbl in header.tables:
        tbl._element.getparent().remove(tbl._element)

    # --- Monta header com tabela: logo | texto ---
    tbl = header.add_table(rows=1, cols=2, width=Inches(6.5))
    tbl.autofit = True

    # Remove bordas da tabela
    tbl_pr = tbl._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl._tbl.insert(0, tbl_pr)

    # Borda none
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tbl_pr.append(borders)

    # Celula 0: Logo
    cell_logo = tbl.cell(0, 0)
    cell_logo.width = Inches(1.4)
    # Clear default paragraph
    cell_logo.paragraphs[0].clear()
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if os.path.exists(LOGO_PATH):
        run_logo = p_logo.add_run()
        run_logo.add_picture(LOGO_PATH, width=Inches(1.25))

    # Celula 1: Texto institucional
    cell_texto = tbl.cell(0, 1)
    cell_texto.width = Inches(5.0)
    # Clear default
    cell_texto.paragraphs[0].clear()

    def add_texto(parag, texto, size, bold, color_hex):
        run = parag.add_run(texto)
        run.font.size = Pt(size)
        run.font.name = 'Times New Roman'
        run.bold = bold
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16),
            int(color_hex[2:4], 16),
            int(color_hex[4:6], 16)
        )
        return run

    p1 = cell_texto.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_texto(p1, "UNIVERSIDADE FEDERAL DO RECÔNCAVO DA BAHIA — UFRB", 8, True, AZUL_UFRB_HEX)

    p2 = cell_texto.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_texto(p2, "CENTRO DE CIÊNCIA E TECNOLOGIA — CETENS", 7.5, True, AZUL_CLARO_HEX)

    p3 = cell_texto.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_texto(p3, "BACHARELADO EM SISTEMAS DE INFORMAÇÃO (EAD)", 7, False, CINZA_HEX)

    p4 = cell_texto.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_texto(p4, "─" * 80, 4, False, DOURADO_HEX)

    # ========== RODAPE ==========
    footer = section.footer
    footer.is_linked_to_previous = False
    # Limpa rodape
    for p in footer.paragraphs:
        p.clear()
    for tbl in footer.tables:
        tbl._element.getparent().remove(tbl._element)

    p_footer = footer.paragraphs[0]
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_footer = p_footer.add_run("MinhaFila — GCETENS843 — 2026.1")
    r_footer.font.size = Pt(7)
    r_footer.font.name = 'Times New Roman'
    r_footer.font.color.rgb = RGBColor(int(CINZA_HEX[0:2], 16), int(CINZA_HEX[2:4], 16), int(CINZA_HEX[4:6], 16))
    r_footer.italic = True

    # ========== ESTILOS ==========
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    for hname, size, chex in [
        ('Heading 1', 14, AZUL_UFRB_HEX),
        ('Heading 2', 13, AZUL_CLARO_HEX),
        ('Heading 3', 12, AZUL_UFRB_HEX),
    ]:
        try:
            h = doc.styles[hname]
            h.font.name = 'Times New Roman'
            h.font.size = Pt(size)
            h.font.bold = True
            h.font.color.rgb = RGBColor(int(chex[0:2], 16), int(chex[2:4], 16), int(chex[4:6], 16))
            h.paragraph_format.line_spacing = 1.5
        except KeyError:
            pass

    doc.save(docx_path)
    print(f"  ✓ molde UFRB aplicado em {os.path.basename(docx_path)}")


def main():
    print(f"🎨 Regenerando .docx com molde UFRB...\n")

    if not os.path.exists(LOGO_PATH):
        print(f"❌ Logo não encontrado: {LOGO_PATH}")
        sys.exit(1)

    for nome in ARQUIVOS:
        print(f"\n📄 {nome}:")
        docx = gerar_docx_via_pandoc(nome)
        if docx:
            aplicar_header_ufrb(docx)

    print(f"\n✅ Pronto! Verifique os arquivos em:")
    for nome in ARQUIVOS:
        path = os.path.join(BASE, f'{nome}.docx')
        size = os.path.getsize(path)
        print(f"   • {nome}.docx ({size/1024:.0f} KB)")


if __name__ == '__main__':
    main()
